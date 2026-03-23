"""Word document extraction using python-docx and docx2txt."""

from pathlib import Path
from typing import Any

from app.core.logging import get_logger
from app.extraction.base import BaseExtractor, ExtractionResult, PageContent

logger = get_logger(__name__)


class DocExtractor(BaseExtractor):
    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    async def extract(self, file_path: Path) -> ExtractionResult:
        try:
            if file_path.suffix.lower() == ".docx":
                return await self._extract_docx(file_path)
            else:
                return await self._extract_doc_fallback(file_path)
        except Exception as e:
            logger.error("Word extraction failed for %s: %s", file_path, e)
            return ExtractionResult(success=False, error=str(e))

    async def _extract_docx(self, file_path: Path) -> ExtractionResult:
        from docx import Document

        doc = Document(str(file_path))
        metadata: dict[str, Any] = {}

        # Extract core properties
        core = doc.core_properties
        metadata["author"] = core.author or ""
        metadata["title"] = core.title or ""
        metadata["subject"] = core.subject or ""
        metadata["created"] = str(core.created) if core.created else ""
        metadata["modified"] = str(core.modified) if core.modified else ""

        # Extract paragraphs
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract tables
        tables_data: list[list[list[str]]] = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables_data.append(table_data)

        metadata["paragraphs_count"] = len(paragraphs)
        metadata["tables_count"] = len(tables_data)
        if tables_data:
            metadata["tables"] = tables_data

        content = "\n\n".join(paragraphs)

        pages = [PageContent(page_number=1, content=content, metadata=metadata)]
        return ExtractionResult(pages=pages, metadata=metadata)

    async def _extract_doc_fallback(self, file_path: Path) -> ExtractionResult:
        """Fallback for .doc files using docx2txt."""
        import docx2txt

        text = docx2txt.process(str(file_path))
        metadata = {"extraction_method": "docx2txt"}

        pages = [PageContent(page_number=1, content=text or "", metadata=metadata)]
        return ExtractionResult(pages=pages, metadata=metadata)
